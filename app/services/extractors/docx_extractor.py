"""
DOCX extraction service.
Extracts text and metadata from Word documents.
"""

import docx
from io import BytesIO
import logging

logger = logging.getLogger(__name__)


def extract_from_docx(file_data):
    """
    Extract data from a DOCX file.
    
    Args:
        file_data: DOCX file as bytes
    
    Returns:
        dict: Extracted text and metadata
    """
    results = {}
    
    try:
        doc = docx.Document(BytesIO(file_data))
        
        # Extract all paragraphs
        all_text = []
        for para in doc.paragraphs:
            if para.text.strip():  # Skip empty paragraphs
                all_text.append(para.text)
        
        if all_text:
            results['extracted_text'] = {
                'text': '\n\n'.join(all_text),
                'paragraph_count': len(all_text),
                'char_count': sum(len(t) for t in all_text)
            }
        
        # Extract tables
        all_tables = []
        for table_num, table in enumerate(doc.tables, 1):
            rows = []
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                rows.append(cells)
            
            if rows:
                all_tables.append({
                    'page': 1,  # DOCX doesn't have page concept
                    'table_number': table_num,
                    'data': rows,
                    'row_count': len(rows),
                    'column_count': len(rows[0]) if rows else 0
                })
        
        if all_tables:
            results['extracted_tables'] = {
                'tables': all_tables,
                'table_count': len(all_tables)
            }
        
        # Extract metadata
        core_props = doc.core_properties
        results['document_metadata'] = {
            'author': core_props.author,
            'title': core_props.title,
            'subject': core_props.subject,
            'created': str(core_props.created) if core_props.created else None,
            'modified': str(core_props.modified) if core_props.modified else None,
            'paragraph_count': len(doc.paragraphs),
            'table_count': len(doc.tables)
        }
        
        logger.info(f"Successfully extracted DOCX: {len(all_text)} paragraphs, {len(all_tables)} tables")
        return results
        
    except Exception as e:
        logger.error(f"Failed to extract DOCX: {e}")
        raise Exception(f"DOCX extraction failed: {str(e)}")